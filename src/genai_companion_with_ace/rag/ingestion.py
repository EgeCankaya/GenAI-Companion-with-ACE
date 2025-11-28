"""Document ingestion pipeline for the IBM Gen AI Companion.

This module is responsible for loading raw course materials (PDFs, Markdown, transcripts,
Jupyter notebooks, etc.), chunking their contents with semantic metadata, and emitting
LangChain document objects that can be passed downstream to the embedding and retrieval
layers. Persisted artifacts are stored under ``data/processed`` so that repeated ingestion
of the same file is avoided.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import uuid
from collections.abc import Iterable, Mapping, MutableMapping, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from genai_companion_with_ace.utils.time import utcnow_isoformat

LOGGER = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".md",
    ".markdown",
    ".txt",
    ".json",
    ".yaml",
    ".yml",
    ".srt",
    ".vtt",
    ".ipynb",
    ".docx",
}


@dataclass(slots=True, frozen=True)
class IngestionConfig:
    """Configuration settings for the document ingestion pipeline."""

    processed_dir: Path = Path("data/processed")
    chunk_size: int = 800
    chunk_overlap: int = 200
    allowed_metadata_keys: Sequence[str] = ("course", "module", "topic", "difficulty", "content_type")
    max_file_size_bytes: int = 10 * 1024 * 1024


@dataclass(slots=True)
class IngestionResult:
    """Summary metadata returned after ingesting a document."""

    document_id: str
    source_path: Path | None
    chunk_count: int
    metadata: dict[str, Any] = field(default_factory=dict)
    artifact_paths: list[Path] = field(default_factory=list)


class DocumentIngestionPipeline:
    """High level orchestrator for ingesting course materials."""

    def __init__(self, config: IngestionConfig | None = None) -> None:
        self._config = config or IngestionConfig()
        self._config.processed_dir.mkdir(parents=True, exist_ok=True)
        self._text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self._config.chunk_size,
            chunk_overlap=self._config.chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )
        LOGGER.debug(
            "Initialized DocumentIngestionPipeline(processed_dir=%s, chunk_size=%s, chunk_overlap=%s)",
            self._config.processed_dir,
            self._config.chunk_size,
            self._config.chunk_overlap,
        )

    @property
    def config(self) -> IngestionConfig:
        """Return the configuration currently in use."""
        return self._config

    def ingest_paths(
        self,
        paths: Iterable[Path],
        *,
        metadata: Mapping[str, Any] | None = None,
        persist: bool = True,
    ) -> list[Document]:
        """Ingest a collection of file paths and return chunked documents."""
        documents: list[Document] = []
        for path in paths:
            docs = self.ingest_path(path, metadata=metadata, persist=persist)
            documents.extend(docs)
        return documents

    def ingest_path(
        self,
        path: Path,
        *,
        metadata: Mapping[str, Any] | None = None,
        persist: bool = True,
    ) -> list[Document]:
        """Ingest a single file system path."""
        self._validate_path(path)
        raw_text = self._load_file(path)
        base_metadata = self._build_base_metadata(source=str(path), metadata=metadata)
        return self._process_raw_text(raw_text, base_metadata=base_metadata, persist=persist)

    def ingest_raw_content(
        self,
        *,
        content: str,
        source_name: str,
        metadata: Mapping[str, Any] | None = None,
        persist: bool = False,
    ) -> list[Document]:
        """Ingest raw string content provided at runtime (e.g., uploaded transcript)."""
        base_metadata = self._build_base_metadata(source=source_name, metadata=metadata)
        return self._process_raw_text(content, base_metadata=base_metadata, persist=persist)

    def _process_raw_text(
        self,
        raw_text: str,
        *,
        base_metadata: MutableMapping[str, Any],
        persist: bool,
    ) -> list[Document]:
        if not raw_text.strip():
            LOGGER.warning("Skipping ingestion for empty document originating from %s", base_metadata.get("source"))
            return []

        chunked_docs = self._text_splitter.create_documents([raw_text], metadatas=[dict(base_metadata)])
        document_id = self._derive_document_id(base_metadata.get("source", "inline"))

        for idx, doc in enumerate(chunked_docs):
            enriched_metadata = doc.metadata
            enriched_metadata["chunk_id"] = idx
            enriched_metadata["document_id"] = document_id
            enriched_metadata.setdefault("ingested_at", utcnow_isoformat())
            doc.page_content = doc.page_content.strip()

        if persist:
            artifact_paths = self._persist_documents(document_id, chunked_docs)
            LOGGER.info(
                "Persisted %s chunks for document %s to %s",
                len(chunked_docs),
                document_id,
                artifact_paths,
            )
        else:
            LOGGER.debug("Skipping persistence for document %s (persist=%s)", document_id, persist)

        return chunked_docs

    def _persist_documents(self, document_id: str, documents: Sequence[Document]) -> list[Path]:
        artifact_dir = self._config.processed_dir / document_id
        artifact_dir.mkdir(parents=True, exist_ok=True)

        records = [
            {
                "document_id": doc.metadata.get("document_id", document_id),
                "chunk_id": doc.metadata.get("chunk_id"),
                "metadata": doc.metadata,
                "content": doc.page_content,
            }
            for doc in documents
        ]

        output_path = artifact_dir / "chunks.jsonl"
        with output_path.open("w", encoding="utf-8") as file:
            for record in records:
                file.write(json.dumps(record, ensure_ascii=False))
                file.write("\n")

        manifest_path = artifact_dir / "manifest.yaml"
        manifest = {
            "document_id": document_id,
            "chunk_count": len(records),
            "metadata": self._filter_metadata(records[0]["metadata"]) if records else {},
        }
        with manifest_path.open("w", encoding="utf-8") as manifest_file:
            yaml.safe_dump(manifest, manifest_file)

        return [output_path, manifest_path]

    @staticmethod
    def _derive_document_id(source: str) -> str:
        file_name = Path(source).name
        digest = hashlib.sha256(source.encode("utf-8")).hexdigest()[:12]
        return f"{file_name}-{digest}"

    def _build_base_metadata(
        self,
        *,
        source: str,
        metadata: Mapping[str, Any] | None,
    ) -> MutableMapping[str, Any]:
        base: MutableMapping[str, Any] = {"source": source}
        if metadata:
            for key, value in metadata.items():
                if key in self._config.allowed_metadata_keys:
                    base[key] = value
                else:
                    LOGGER.debug("Ignoring unsupported metadata key '%s' for source %s", key, source)
        return base

    def _validate_path(self, path: Path) -> None:
        if not path.exists():
            raise FileNotFoundError(path)
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            message = f"Unsupported file extension '{path.suffix}'. Supported extensions: {allowed}"
            raise ValueError(message)
        if self._config.max_file_size_bytes > 0 and path.stat().st_size > self._config.max_file_size_bytes:
            message = (
                f"{path.name} exceeds the configured ingestion limit of "
                f"{self._config.max_file_size_bytes / (1024 * 1024):.1f} MB"
            )
            raise ValueError(message)

    def _load_file(self, path: Path) -> str:
        loader_map = {
            ".pdf": self._load_pdf,
            ".md": self._load_text,
            ".markdown": self._load_text,
            ".txt": self._load_text,
            ".json": self._load_json,
            ".yaml": self._load_yaml,
            ".yml": self._load_yaml,
            ".srt": self._load_text,
            ".vtt": self._load_text,
            ".ipynb": self._load_notebook,
            ".docx": self._load_docx,
        }
        loader = loader_map.get(path.suffix.lower())
        if loader is None:
            message = f"No loader registered for extension: {path.suffix}"
            raise ValueError(message)
        LOGGER.debug("Loading document %s using loader %s", path, loader.__name__)
        return loader(path)

    @staticmethod
    def _load_text(path: Path) -> str:
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _load_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - dependency should exist
            message = "Install 'pypdf' to enable PDF ingestion"
            raise RuntimeError(message) from exc

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    @staticmethod
    def _load_docx(path: Path) -> str:
        try:
            import docx  # type: ignore[import]
        except ImportError as exc:  # pragma: no cover
            message = "Install 'python-docx' to enable DOCX ingestion"
            raise RuntimeError(message) from exc

        document = docx.Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(paragraphs)

    @staticmethod
    def _load_json(path: Path) -> str:
        with path.open("r", encoding="utf-8") as file:
            content = json.load(file)
        return json.dumps(content, indent=2, ensure_ascii=False)

    @staticmethod
    def _load_yaml(path: Path) -> str:
        with path.open("r", encoding="utf-8") as file:
            content = yaml.safe_load(file)
        return yaml.safe_dump(content, sort_keys=False, allow_unicode=True)

    @staticmethod
    def _load_notebook(path: Path) -> str:
        with path.open("r", encoding="utf-8") as file:
            notebook = json.load(file)

        cells = notebook.get("cells", [])
        buffers: list[str] = []
        for cell in cells:
            cell_type = cell.get("cell_type")
            source = "".join(cell.get("source", []))
            if cell_type == "markdown":
                buffers.append(source)
            elif cell_type == "code":
                buffers.append(f"```python\n{source}\n```")
        return "\n\n".join(buffers)

    def create_manifest(
        self,
        *,
        document_id: str | None = None,
        metadata: Mapping[str, Any] | None = None,
        chunk_count: int | None = None,
        artifact_paths: Sequence[Path] | None = None,
    ) -> IngestionResult:
        """Create a synthetic manifest entry for programmatic ingestion results."""
        document_id = document_id or uuid.uuid4().hex
        artifact_paths_list = list(artifact_paths or [])
        result = IngestionResult(
            document_id=document_id,
            source_path=None,
            chunk_count=chunk_count or 0,
            metadata=dict(metadata or {}),
            artifact_paths=artifact_paths_list,
        )
        manifest_path = self._config.processed_dir / document_id / "manifest.yaml"
        manifest_path.parent.mkdir(parents=True, exist_ok=True)
        with manifest_path.open("w", encoding="utf-8") as manifest_file:
            yaml.safe_dump(
                {
                    "document_id": result.document_id,
                    "metadata": self._filter_metadata(result.metadata),
                    "chunk_count": result.chunk_count,
                    "artifact_paths": [str(path) for path in artifact_paths_list],
                },
                manifest_file,
            )
        result.artifact_paths.append(manifest_path)
        return result

    def list_ingested_documents(self) -> list[IngestionResult]:
        """Return summaries for all previously ingested documents."""
        results: list[IngestionResult] = []
        for manifest_path in self._config.processed_dir.glob("*/manifest.yaml"):
            with manifest_path.open("r", encoding="utf-8") as manifest_file:
                payload = yaml.safe_load(manifest_file) or {}
            artifact_dir = manifest_path.parent
            chunks_path = artifact_dir / "chunks.jsonl"
            results.append(
                IngestionResult(
                    document_id=payload.get("document_id", artifact_dir.name),
                    source_path=None,
                    chunk_count=payload.get("chunk_count", 0),
                    metadata=payload.get("metadata", {}),
                    artifact_paths=[chunks_path, manifest_path] if chunks_path.exists() else [manifest_path],
                )
            )
        return sorted(results, key=lambda result: result.document_id)

    def _filter_metadata(self, metadata: Mapping[str, Any]) -> dict[str, Any]:
        return {key: metadata[key] for key in metadata if key in self._config.allowed_metadata_keys}


def load_processed_document(document_dir: Path) -> list[Document]:
    """Utility function to load persisted chunks from disk back into Document objects."""
    chunks_path = document_dir / "chunks.jsonl"
    if not chunks_path.exists():
        raise FileNotFoundError(chunks_path)

    documents: list[Document] = []
    with chunks_path.open("r", encoding="utf-8") as file:
        for line in file:
            payload = json.loads(line)
            documents.append(Document(page_content=payload["content"], metadata=payload["metadata"]))
    return documents


def infer_metadata_from_path(path: Path) -> dict[str, Any]:
    """Heuristic helper to infer course/module metadata from a file path."""
    parts = [part for part in path.parts if part not in {"data", "raw", "processed"}]
    metadata: dict[str, Any] = {"content_type": path.suffix.lstrip(".")}
    if parts:
        metadata["source_hint"] = os.path.join(*parts[:-1]) if len(parts) > 1 else path.parent.name
    return metadata

